"""Create test fixtures for academic translation testing."""

import os
from pathlib import Path

from docx import Document


def create_fixtures_dir() -> Path:
    """Create fixtures directory."""
    fixtures_dir = Path(__file__).parent
    fixtures_dir.mkdir(exist_ok=True)
    return fixtures_dir


def create_minimal_docx() -> Path:
    """Create paper_minimal.docx with headings, citations, and numbers."""
    fixtures_dir = create_fixtures_dir()
    output_path = fixtures_dir / "paper_minimal.docx"

    doc = Document()

    # Title
    title = doc.add_paragraph("Minimal Academic Paper")
    title.style = "Heading 1"

    # Abstract
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This paper presents a minimal academic document for testing. "
        "We demonstrate key features including citations (Smith, 2023) "
        "and numerical values (p < 0.05)."
    )

    # Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Academic research relies on proper citation formatting. "
        "Studies [1, 2, 3] have shown that accurate citations are essential. "
        "The value of α was set to 0.95 with a confidence interval of 95%."
    )

    # Methodology
    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(
        "We used a sample size of N=100 participants. The results show "
        "statistical significance with R² = 0.87 and an effect size of d = 0.5."
    )

    # References
    doc.add_heading("References", level=1)
    doc.add_paragraph("Smith, J. (2023). Citation practices in academic writing.")
    doc.add_paragraph("Johnson, A. et al. (2022). Statistical methods for research.")
    doc.add_paragraph("Brown, R. et al. (2021). Effect size calculations.")

    doc.save(output_path)
    return output_path


def create_citations_docx() -> Path:
    """Create paper_citations.docx with varied citation formats."""
    fixtures_dir = create_fixtures_dir()
    output_path = fixtures_dir / "paper_citations.docx"

    doc = Document()

    # Title
    title = doc.add_paragraph("Citation Format Examples")
    title.style = "Heading 1"

    # Various citation formats
    doc.add_heading("Citation Formats", level=1)

    doc.add_paragraph(
        "Parenthetical citations: (Smith, 2020), (Johnson and Williams, 2021)."
    )

    doc.add_paragraph(
        "Et al. citations: (Brown et al., 2022), (Davis et al. 2023)."
    )

    doc.add_paragraph("Bracked citations: [1], [2, 3], [4, 5, 6].")

    doc.add_paragraph(
        "Mixed format: Recent work (Lee et al., 2022) [7] demonstrates this."
    )

    doc.add_paragraph(
        "Citations with page numbers: (Smith, 2020, p. 45), (Johnson, 2021, pp. 100-102)."
    )

    # References
    doc.add_heading("References", level=1)
    doc.add_paragraph("Smith, J. (2020). First citation example.")
    doc.add_paragraph("Johnson, A., & Williams, B. (2021). Second citation example.")
    doc.add_paragraph("Brown, C., et al. (2022). Et al. citation example.")
    doc.add_paragraph("Davis, E., et al. (2023). Another et al. example.")
    doc.add_paragraph("Lee, F., et al. (2022). Mixed format citation.")

    doc.save(output_path)
    return output_path


def create_numbers_docx() -> Path:
    """Create paper_numbers.docx with p-values, ranges, decimals, percents."""
    fixtures_dir = create_fixtures_dir()
    output_path = fixtures_dir / "paper_numbers.docx"

    doc = Document()

    # Title
    title = doc.add_paragraph("Numerical Data Examples")
    title.style = "Heading 1"

    # Numbers
    doc.add_heading("Statistical Results", level=1)

    doc.add_paragraph(
        "P-values: The results were significant (p < 0.05), (p = 0.01), "
        "and (p < 0.001). Non-significant: (p > 0.05), (p = 0.15)."
    )

    doc.add_paragraph(
        "Ranges: Age range was 18-65 years. The 95% CI was [0.45, 0.89]. "
        "Temperature range: 20°C to 25°C."
    )

    doc.add_paragraph(
        "Decimals: The mean was 3.14159 with SD = 0.78234. "
        "Precision: 0.0001, 1.2345e-5, 6.022e23."
    )

    doc.add_paragraph(
        "Percentages: Response rate was 87.5%. The sample was 45.3% female. "
        "Improvement: +23.4%, -5.2%."
    )

    doc.add_paragraph(
        "Regression: R² = 0.8765, adjusted R² = 0.8321. "
        "F-statistic: F(2, 97) = 45.32."
    )

    doc.add_paragraph(
        "Sample sizes: N = 150, n = 30 per group, total Ntotal = 300."
    )

    # Results
    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        "Table 1 shows the results. Mean difference: 2.35 (SD = 1.12). "
        "Effect size: Cohen's d = 0.68, 95% CI [0.12, 1.24]."
    )

    doc.save(output_path)
    return output_path


def create_references_docx() -> Path:
    """Create paper_references.docx for reference section detection."""
    fixtures_dir = create_fixtures_dir()
    output_path = fixtures_dir / "paper_references.docx"

    doc = Document()

    # Title
    title = doc.add_paragraph("Reference Section Detection Test")
    title.style = "Heading 1"

    # Main content
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This is an example paper for testing reference section detection. "
        "We cite several works (Smith, 2020), (Johnson et al., 2021), and [1, 2]."
    )

    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "Our approach is based on previous research. See (Brown, 2022) for details."
    )

    # References - using different heading styles
    doc.add_heading("References", level=1)
    doc.add_paragraph("Smith, J. (2020). First reference entry.")
    doc.add_paragraph("Johnson, A., Brown, B., & Davis, C. (2021). Second reference entry.")
    doc.add_paragraph("Brown, D. (2022). Third reference entry.")

    # Test "Bibliography" alternative
    doc.add_heading("Bibliography", level=1)
    doc.add_paragraph("Lee, E. (2023). Bibliography entry 1.")
    doc.add_paragraph("Williams, F. (2023). Bibliography entry 2.")

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    print("Creating test fixtures...")

    paths = [
        create_minimal_docx(),
        create_citations_docx(),
        create_numbers_docx(),
        create_references_docx(),
    ]

    for path in paths:
        if path.exists():
            size = path.stat().st_size
            print(f"  Created: {path.name} ({size} bytes)")
        else:
            print(f"  Failed: {path.name}")

    print("Done!")
