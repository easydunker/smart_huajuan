"""Tests for DOCX export (M6 Phase 2)."""

import pytest
from docx import Document as DocxDocument

from aat.export.docx_export import DocxExporter
from aat.export.global_pass import GlobalPassReport
from aat.storage.models import (
    DocumentModel,
    Reference,
    Segment,
    SegmentState,
    TranslationProject,
    TranslationSegment,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_seg(
    sid: str,
    source: str,
    translation: str,
    locked: bool = True,
    chapter_id: str | None = None,
) -> TranslationSegment:
    seg = Segment(
        sid=sid,
        pid_list=[],
        source_text=source,
        chapter_id=chapter_id,
    )
    return TranslationSegment(
        segment=seg,
        state=SegmentState.LOCK_SEGMENT,
        translation=translation,
        locked=locked,
    )


def _make_project(
    segments: list[TranslationSegment],
    references: list[Reference] | None = None,
    title: str = "Test Document",
) -> TranslationProject:
    doc = DocumentModel(
        doc_id="test-doc",
        title=title,
        sections=[],
        references=references or [],
        citations=[],
    )
    proj = TranslationProject(project_id="test-proj", document=doc)
    proj.segments = segments
    return proj


# ===================================================================
# TestDocxExporter
# ===================================================================


class TestDocxExporter:
    """Tests for DocxExporter."""

    def test_export_produces_valid_docx(self, tmp_path):
        """Export produces a valid .docx readable by python-docx."""
        segs = [_make_seg("s1", "Hello world.", "你好世界。", chapter_id="Chapter 1")]
        project = _make_project(segs)
        exporter = DocxExporter(project)
        out = exporter.export(tmp_path / "out.docx")

        assert out.exists()
        doc = DocxDocument(str(out))
        assert len(doc.paragraphs) > 0

    def test_chapter_headings_present(self, tmp_path):
        """Chapter headings appear as Heading 1 in the document."""
        segs = [
            _make_seg("s1", "Intro text.", "介绍文本。", chapter_id="Introduction"),
            _make_seg("s2", "Method text.", "方法文本。", chapter_id="Methods"),
        ]
        project = _make_project(segs)
        exporter = DocxExporter(project)
        out = exporter.export(tmp_path / "out.docx")

        doc = DocxDocument(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert "Introduction" in headings
        assert "Methods" in headings

    def test_citations_preserved_in_text(self, tmp_path):
        """Citations like (Smith, 2020) appear in paragraph text."""
        segs = [
            _make_seg(
                "s1",
                "As (Smith, 2020) notes, this is important.",
                "正如(Smith, 2020)所指出的，这很重要。",
                chapter_id="Chapter 1",
            ),
        ]
        project = _make_project(segs)
        exporter = DocxExporter(project)
        out = exporter.export(tmp_path / "out.docx")

        doc = DocxDocument(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "(Smith, 2020)" in all_text

    def test_bilingual_mode_interleaved(self, tmp_path):
        """Bilingual mode: source italic + translation normal + separator."""
        source = "Machine learning is powerful."
        translation = "机器学习很强大。"
        segs = [_make_seg("s1", source, translation, chapter_id="Ch1")]
        project = _make_project(segs)
        exporter = DocxExporter(project, bilingual=True)
        out = exporter.export(tmp_path / "out.docx")

        doc = DocxDocument(str(out))

        body_paras = []
        found_page_break = False
        for p in doc.paragraphs:
            if found_page_break:
                body_paras.append(p)
            if any(
                run.text == "" and run._element.xml.find("lastRenderedPageBreak") != -1
                for run in p.runs
            ):
                found_page_break = True
            if not found_page_break:
                for run in p.runs:
                    if hasattr(run._element, "xml") and "w:br" in run._element.xml and 'w:type="page"' in run._element.xml:
                        found_page_break = True
                        break

        source_paras = [p for p in doc.paragraphs if p.text == source]
        translation_paras = [p for p in doc.paragraphs if p.text == translation]
        separator_paras = [p for p in doc.paragraphs if "─" in p.text]

        assert len(source_paras) >= 1, "Source text paragraph not found"
        assert len(translation_paras) >= 1, "Translation paragraph not found"
        assert len(separator_paras) >= 1, "Separator paragraph not found"

        src_para = source_paras[0]
        assert src_para.runs[0].italic is True, "Source text should be italic"
        src_color = src_para.runs[0].font.color.rgb
        assert src_color is not None, "Source text should have grey colour"

    def test_metadata_page_present(self, tmp_path):
        """Metadata page contains title, model info, and segment count."""
        segs = [
            _make_seg("s1", "Text.", "文本。", chapter_id="Ch1"),
            _make_seg("s2", "More.", "更多。", chapter_id="Ch1"),
        ]
        project = _make_project(segs, title="My Dissertation")
        exporter = DocxExporter(
            project, model_provider="anthropic", model_name="claude-3"
        )
        out = exporter.export(tmp_path / "out.docx")

        doc = DocxDocument(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "My Dissertation" in all_text
        assert "anthropic" in all_text
        assert "claude-3" in all_text
        assert "Total segments: 2" in all_text

    def test_metadata_page_global_pass_status(self, tmp_path):
        """Metadata page shows global pass status when report is provided."""
        segs = [_make_seg("s1", "Text.", "文本。", chapter_id="Ch1")]
        project = _make_project(segs)
        report = GlobalPassReport(passed=True, summary="All checks passed")
        exporter = DocxExporter(project, global_report=report)
        out = exporter.export(tmp_path / "out.docx")

        doc = DocxDocument(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "PASSED" in all_text

    def test_references_section_included(self, tmp_path):
        """References section appears when the document has references."""
        refs = [
            Reference(rid="r1", raw="Smith, J. (2020). A study. Journal, 1(1), 1-10."),
            Reference(rid="r2", raw="Jones, K. (2019). Another study. Journal, 2(2), 20-30."),
        ]
        segs = [_make_seg("s1", "Text.", "文本。", chapter_id="Ch1")]
        project = _make_project(segs, references=refs)
        exporter = DocxExporter(project)
        out = exporter.export(tmp_path / "out.docx")

        doc = DocxDocument(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert "References" in headings

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Smith, J. (2020)" in all_text
        assert "Jones, K. (2019)" in all_text

    def test_empty_project_produces_valid_docx(self, tmp_path):
        """Empty project (no segments) still produces a valid .docx."""
        project = _make_project([])
        exporter = DocxExporter(project)
        out = exporter.export(tmp_path / "out.docx")

        assert out.exists()
        doc = DocxDocument(str(out))
        assert len(doc.paragraphs) > 0
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Total segments: 0" in all_text

    def test_unlocked_segments_excluded(self, tmp_path):
        """Unlocked segments do not appear in the exported document body."""
        segs = [
            _make_seg("s1", "Locked text.", "锁定文本。", locked=True, chapter_id="Ch1"),
            _make_seg("s2", "Unlocked draft.", "未锁定草稿。", locked=False, chapter_id="Ch1"),
        ]
        project = _make_project(segs)
        exporter = DocxExporter(project)
        out = exporter.export(tmp_path / "out.docx")

        doc = DocxDocument(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "锁定文本。" in all_text
        assert "未锁定草稿。" not in all_text
