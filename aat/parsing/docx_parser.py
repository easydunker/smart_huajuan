"""DOCX parser for extracting document structure and content."""

from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING
from uuid import uuid4

from docx import Document

from aat.storage.models import Citation, DocumentModel, Paragraph, Reference, Section

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument


class DocxParserError(Exception):
    """Exception raised for parsing errors."""


@dataclass
class CitationMatch:
    """Represents a matched citation."""
    text: str
    start: int
    end: int


class DocxParser:
    """Parser for extracting structure and content from DOCX documents."""

    def __init__(self) -> None:
        """Initialize the DOCX parser."""
        self._citations: list[Citation] = []
        self._references: list[Reference] = []

    def parse(self, file_path: str | Path) -> DocumentModel:
        """
        Parse a DOCX file and extract document structure.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            DocumentModel containing parsed document structure.

        Raises:
            DocxParserError: If parsing fails.
            FileNotFoundError: If file doesn't exist.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        try:
            doc = Document(str(path))
        except Exception as e:
            raise DocxParserError(f"Failed to read DOCX file: {e}") from e

        # Initialize document model
        model = DocumentModel.create()
        model.title = self._extract_title(doc)

        # Parse document sections
        sections = self._parse_sections(doc)
        model.sections = sections

        # Store citations and references
        model.citations = self._citations
        model.references = self._references

        return model

    def _extract_title(self, doc: "DocxDocument") -> str | None:
        """
        Extract the document title.

        The title is assumed to be the first paragraph with heading style
        or the first paragraph if no heading is found.

        Args:
            doc: python-docx Document object.

        Returns:
            Extracted title or None if document is empty.
        """
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it has a heading style
                if para.style.name.startswith("Heading 1"):
                    return para.text.strip()
                # Return first non-empty paragraph as fallback
                if len(para.text.strip()) > 0 and len(doc.paragraphs) > 1:
                    # Only use first para if it's short and likely a title
                    if len(para.text.strip()) < 200:
                        return para.text.strip()
        return None

    def _parse_sections(self, doc: "DocxDocument") -> list[Section]:
        """
        Parse document into sections.

        A section starts with a heading and contains all paragraphs until
        the next heading of same or higher level.

        Args:
            doc: python-docx Document object.

        Returns:
            List of Section objects.
        """
        sections: list[Section] = []
        current_heading: str | None = None
        current_paragraphs: list[Paragraph] = []
        current_section_level = 0

        # First, collect all paragraph references
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a heading
            heading_level, heading_text = self._get_heading_info(para)

            if heading_level is not None:
                # Save current section if it exists
                if current_paragraphs or current_heading:
                    sections.append(
                        Section(heading=current_heading, paragraphs=current_paragraphs)
                    )

                # Start new section
                current_heading = heading_text
                current_paragraphs = []
                current_section_level = heading_level
            else:
                # Add paragraph to current section
                pid = str(uuid4())
                citations = self._extract_citations(text, pid)
                paragraph = Paragraph(pid=pid, text=text, citations=citations)
                current_paragraphs.append(paragraph)

        # Save the last section
        if current_paragraphs or current_heading:
            sections.append(Section(heading=current_heading, paragraphs=current_paragraphs))

        # Detect and separate references section
        sections = self._separate_references_section(sections)

        return sections

    def _get_heading_info(self, para) -> tuple[int | None, str | None]:
        """
        Get heading level and text if paragraph is a heading.

        Args:
            para: python-docx Paragraph object.

        Returns:
            Tuple of (level, text) where level is 1-6 or None.
        """
        style_name = para.style.name

        # Match heading patterns
        for level in range(1, 7):
            if style_name == f"Heading {level}":
                return level, para.text.strip()

        # Try to match "Heading X" pattern with numbers
        import re

        match = re.match(r"Heading\s+(\d)", style_name, re.IGNORECASE)
        if match:
            level = int(match.group(1))
            return level, para.text.strip()

        # Check for Title style
        if style_name.lower() == "title":
            return 1, para.text.strip()

        return None, None

    def _extract_citations(self, text: str, pid: str) -> list[Citation]:
        """
        Extract citations from paragraph text.

        Supports patterns:
        - (Author, 2020)
        - (Author et al., 2020)
        - [12]
        - (Smith et al. 2020)

        Args:
            text: Paragraph text.
            pid: Paragraph ID.

        Returns:
            List of Citation objects.
        """
        import re

        citations = []
        citation_patterns = [
            # Parentheses with year: (Smith, 2020) or (Smith et al., 2020)
            r"\([A-Z][a-z]+(?:\s+et\s+al\.?)?(?:,\s+\d{4})?\)",
            # Square brackets: [12] or [12, 13]
            r"\[\d+(?:,\s*\d+)*\]",
            # Parentheses without comma: (Smith et al. 2020)
            r"\([A-Z][a-z]+(?:\s+et\s+al\.?)?\s+\d{4}\)",
        ]

        for pattern in citation_patterns:
            for match in re.finditer(pattern, text):
                citation_text = match.group()
                cid = str(uuid4())
                citation = Citation(cid=cid, text=citation_text, pid=pid)
                citations.append(citation)
                self._citations.append(citation)

        return citations

    def _separate_references_section(self, sections: list[Section]) -> list[Section]:
        """
        Detect and extract references section.

        Looks for sections with heading containing "reference" or "bibliography".

        Args:
            sections: List of all sections.

        Returns:
            Tuple of (main_sections, references_section).
        """
        main_sections = []
        ref_section_found = False

        for section in sections:
            heading_lower = (section.heading or "").lower()

            # Check if this is a references section
            if not ref_section_found and any(
                keyword in heading_lower
                for keyword in ["reference", "bibliography", "works cited"]
            ):
                # Extract references from this and following sections
                self._extract_references(section.paragraphs)
                ref_section_found = True
                # Don't add references section to main sections
            elif ref_section_found:
                # All sections after references are also reference sections
                self._extract_references(section.paragraphs)
            else:
                main_sections.append(section)

        return main_sections

    def _extract_references(self, paragraphs: list[Paragraph]) -> None:
        """
        Extract reference entries from paragraphs.

        Args:
            paragraphs: List of paragraphs in references section.
        """
        for para in paragraphs:
            # Each paragraph is typically one reference entry
            # (or may be split across multiple paragraphs)
            if para.text.strip():
                rid = str(uuid4())
                ref = Reference(rid=rid, raw=para.text.strip())
                self._references.append(ref)
