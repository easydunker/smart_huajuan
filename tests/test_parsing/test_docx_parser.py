"""Tests for DocxParser."""

import os
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

from aat.parsing.docx_parser import DocxParser, DocxParserError
from aat.storage.models import Citation


@pytest.fixture
def sample_docx_path() -> str:
    """Create a sample DOCX file for testing."""
    # Create a temporary DOCX file
    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    temp_path = temp_file.name
    temp_file.close()

    doc = Document()

    # Add title
    title = doc.add_paragraph("Sample Academic Paper")
    title.style = "Heading 1"

    # Add abstract section
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This paper presents a novel approach to machine learning. "
        "Our method achieves state-of-the-art results (Smith et al., 2023)."
    )

    # Add introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Machine learning has revolutionized many fields [1, 2].")
    doc.add_paragraph(
        "Recent work by Johnson et al. (2022) demonstrates the importance of "
        "proper data preprocessing."
    )

    # Add method section
    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(
        "We propose a new algorithm based on neural networks. "
        "The approach combines ideas from deep learning (Brown, 2021) "
        "and classical statistics."
    )

    # Add results
    doc.add_heading("Results", level=1)
    doc.add_paragraph("Our experiments show 95% accuracy with p < 0.05.")

    # Add references
    doc.add_heading("References", level=1)
    doc.add_paragraph("Smith, J., et al. (2023). Novel machine learning approach.")
    doc.add_paragraph("Johnson, A., et al. (2022). Data preprocessing for ML.")
    doc.add_paragraph("Brown, R. (2021). Neural network fundamentals.")

    doc.save(temp_path)
    yield temp_path

    # Cleanup
    if os.path.exists(temp_path):
        os.unlink(temp_path)


class TestDocxParser:
    """Test DocxParser functionality."""

    def test_parse_nonexistent_file(self) -> None:
        """Test parsing a non-existent file."""
        parser = DocxParser()
        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/file.docx")

    def test_parse_basic_document(self, sample_docx_path: str) -> None:
        """Test parsing a basic document."""
        parser = DocxParser()
        model = parser.parse(sample_docx_path)

        assert model.doc_id
        assert model.title == "Sample Academic Paper"
        assert len(model.sections) > 0

    def test_extract_title(self, sample_docx_path: str) -> None:
        """Test title extraction."""
        parser = DocxParser()
        model = parser.parse(sample_docx_path)

        assert model.title == "Sample Academic Paper"

    def test_extract_sections(self, sample_docx_path: str) -> None:
        """Test section extraction."""
        parser = DocxParser()
        model = parser.parse(sample_docx_path)

        # Should have sections: Abstract, Introduction, Methodology, Results
        # (References should be separated)
        section_headings = [s.heading for s in model.sections if s.heading]
        assert "Abstract" in section_headings
        assert "Introduction" in section_headings
        assert "Methodology" in section_headings
        assert "Results" in section_headings

        # References should not be in main sections
        assert "References" not in section_headings

    def test_extract_citations(self, sample_docx_path: str) -> None:
        """Test citation extraction."""
        parser = DocxParser()
        model = parser.parse(sample_docx_path)

        # Should extract citations from the document
        assert len(model.citations) > 0

        # Check citation types
        citation_texts = [c.text for c in model.citations]

        # Check for different citation formats
        parenthetical = any("(" in text and ")" in text for text in citation_texts)
        bracketed = any("[" in text and "]" in text for text in citation_texts)

        assert parenthetical or bracketed

    def test_extract_references(self, sample_docx_path: str) -> None:
        """Test reference extraction."""
        parser = DocxParser()
        model = parser.parse(sample_docx_path)

        # Should have extracted references
        assert len(model.references) == 3

        # Check reference content
        ref_texts = [r.raw for r in model.references]
        assert any("Smith" in text for text in ref_texts)
        assert any("Johnson" in text for text in ref_texts)
        assert any("Brown" in text for text in ref_texts)

    def test_paragraphs_have_ids(self, sample_docx_path: str) -> None:
        """Test that paragraphs have unique IDs."""
        parser = DocxParser()
        model = parser.parse(sample_docx_path)

        all_pids = []
        for section in model.sections:
            for para in section.paragraphs:
                all_pids.append(para.pid)

        # All paragraphs should have IDs
        assert len(all_pids) > 0
        assert all(pid for pid in all_pids)

        # IDs should be unique
        assert len(all_pids) == len(set(all_pids))

    def test_citations_have_paragraph_pointers(self, sample_docx_path: str) -> None:
        """Test that citations point to correct paragraphs."""
        parser = DocxParser()
        model = parser.parse(sample_docx_path)

        # Collect all valid pids
        all_pids = set()
        for section in model.sections:
            for para in section.paragraphs:
                all_pids.add(para.pid)

        # All citations should point to valid paragraphs
        for citation in model.citations:
            assert citation.pid in all_pids


class TestCitationExtraction:
    """Test citation extraction patterns."""

    def test_citation_pattern_parenthetical(self) -> None:
        """Test (Author, 2020) pattern."""
        temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        temp_path = temp_file.name
        temp_file.close()

        doc = Document()
        doc.add_paragraph("This is a citation (Smith, 2020) in text.")
        doc.save(temp_path)

        parser = DocxParser()
        model = parser.parse(temp_path)

        citation_texts = [c.text for c in model.citations]
        assert "(Smith, 2020)" in citation_texts

        os.unlink(temp_path)

    def test_citation_pattern_et_al(self) -> None:
        """Test (Author et al., 2020) pattern."""
        temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        temp_path = temp_file.name
        temp_file.close()

        doc = Document()
        doc.add_paragraph("This work (Smith et al., 2020) shows results.")
        doc.save(temp_path)

        parser = DocxParser()
        model = parser.parse(temp_path)

        citation_texts = [c.text for c in model.citations]
        assert "(Smith et al., 2020)" in citation_texts

        os.unlink(temp_path)

    def test_citation_pattern_bracketed(self) -> None:
        """Test [12] pattern."""
        temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        temp_path = temp_file.name
        temp_file.close()

        doc = Document()
        doc.add_paragraph("Multiple studies [1, 2, 3] support this.")
        doc.save(temp_path)

        parser = DocxParser()
        model = parser.parse(temp_path)

        citation_texts = [c.text for c in model.citations]
        assert any("[1, 2, 3]" in text for text in citation_texts)

        os.unlink(temp_path)

    def test_citation_without_comma(self) -> None:
        """Test (Author et al. 2020) pattern without comma."""
        temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        temp_path = temp_file.name
        temp_file.close()

        doc = Document()
        doc.add_paragraph("Recent work (Brown et al. 2021) demonstrates this.")
        doc.save(temp_path)

        parser = DocxParser()
        model = parser.parse(temp_path)

        citation_texts = [c.text for c in model.citations]
        assert "(Brown et al. 2021)" in citation_texts

        os.unlink(temp_path)
