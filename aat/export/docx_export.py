"""DOCX export with bilingual mode and metadata page."""

from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

if TYPE_CHECKING:
    from aat.export.global_pass import GlobalPassReport
    from aat.storage.models import TranslationProject, TranslationSegment


class DocxExportError(Exception):
    """Exception raised for DOCX export errors."""


class DocxExporter:
    """Produce a ``.docx`` with proper formatting from a TranslationProject."""

    def __init__(
        self,
        project: "TranslationProject",
        *,
        bilingual: bool = False,
        global_report: "GlobalPassReport | None" = None,
        model_provider: str = "unknown",
        model_name: str = "unknown",
    ) -> None:
        self._project = project
        self._bilingual = bilingual
        self._report = global_report
        self._model_provider = model_provider
        self._model_name = model_name

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def export(self, output_path: str | Path) -> Path:
        """Build and save a .docx file.  Returns the output path."""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._add_metadata_page(doc)
        self._add_body(doc)
        try:
            doc.save(str(output_path))
        except Exception as e:
            raise DocxExportError(str(e)) from e
        return output_path

    # ------------------------------------------------------------------
    # Internals
    # ------------------------------------------------------------------

    def _add_metadata_page(self, doc: Document) -> None:
        """First page: translation metadata."""
        title = self._project.document.title or self._project.document.doc_id
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(20)

        doc.add_paragraph("")  # spacer

        meta_lines = [
            f"Source document: {self._project.document.doc_id}",
            f"Translation date: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            f"Model: {self._model_provider} / {self._model_name}",
            f"Total segments: {len(self._project.segments)}",
        ]

        if self._report is not None:
            status = "PASSED" if self._report.passed else "ISSUES FOUND"
            meta_lines.append(f"Global pass: {status}")
            if not self._report.passed:
                meta_lines.append(f"  {self._report.summary}")

        for line in meta_lines:
            doc.add_paragraph(line, style="Normal")

        doc.add_page_break()

    def _add_body(self, doc: Document) -> None:
        """Add translated content grouped by chapter."""
        segments = self._project.segments
        locked = [s for s in segments if s.locked and s.translation]

        chapter_groups = self._group_by_chapter(locked)

        for chapter_id, chapter_segs in chapter_groups:
            if chapter_id:
                doc.add_heading(chapter_id, level=1)

            for seg in chapter_segs:
                if self._bilingual:
                    self._add_bilingual_segment(doc, seg)
                else:
                    doc.add_paragraph(seg.translation or "", style="Normal")

        # Reference section pass-through
        refs = self._project.document.references
        if refs:
            doc.add_heading("References", level=1)
            for ref in refs:
                doc.add_paragraph(ref.raw, style="Normal")

    def _add_bilingual_segment(
        self, doc: Document, seg: "TranslationSegment"
    ) -> None:
        """Add source (grey italic) then translation (normal)."""
        # Source
        src_para = doc.add_paragraph()
        src_run = src_para.add_run(seg.segment.source_text)
        src_run.italic = True
        src_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Translation
        doc.add_paragraph(seg.translation or "", style="Normal")

        # Separator
        sep = doc.add_paragraph()
        sep_run = sep.add_run("─" * 40)
        sep_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    @staticmethod
    def _group_by_chapter(
        segments: list["TranslationSegment"],
    ) -> list[tuple[str | None, list["TranslationSegment"]]]:
        """Group segments by chapter_id preserving order."""
        groups: dict[str | None, list["TranslationSegment"]] = {}
        for seg in segments:
            ch = seg.segment.chapter_id
            groups.setdefault(ch, []).append(seg)
        return list(groups.items())
