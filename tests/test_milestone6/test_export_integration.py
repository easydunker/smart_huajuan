"""End-to-end integration tests for global pass + DOCX export (M6)."""

import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from aat.export.docx_export import DocxExporter
from aat.export.global_pass import GlobalPassOrchestrator
from aat.storage.models import (
    DocumentModel,
    Reference,
    Segment,
    SegmentState,
    TranslationProject,
    TranslationSegment,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _seg(
    sid: str,
    source: str,
    translation: str,
    chapter_id: str = "chapter1",
) -> TranslationSegment:
    """Create a locked TranslationSegment."""
    return TranslationSegment(
        segment=Segment(
            sid=sid,
            pid_list=[],
            source_text=source,
            chapter_id=chapter_id,
        ),
        state=SegmentState.LOCK_SEGMENT,
        translation=translation,
        locked=True,
    )


def _project_with_segments(
    segments: list[TranslationSegment],
    references: list[Reference] | None = None,
) -> TranslationProject:
    doc = DocumentModel(
        doc_id="integ-doc",
        title="Integration Test Document",
        sections=[],
        references=references or [],
        citations=[],
    )
    proj = TranslationProject(project_id="integ-proj", document=doc)
    proj.segments = segments
    return proj


# ===================================================================
# Integration tests
# ===================================================================


class TestExportIntegration:
    """End-to-end: create project -> global pass -> DOCX export."""

    def test_end_to_end_export(self, tmp_path: Path) -> None:
        """Full pipeline: 5 segments across 2 chapters, global pass, DOCX."""
        segments = [
            _seg("s1", "Introduction to the study.", "本研究的引言。", "chapter1"),
            _seg("s2", "Methods are described below.", "方法如下所述。", "chapter1"),
            _seg("s3", "Results show improvement.", "结果显示有所改善。", "chapter2"),
            _seg("s4", "Discussion of findings.", "研究发现的讨论。", "chapter2"),
            _seg("s5", "Conclusion and future work.", "结论与未来工作。", "chapter2"),
        ]
        project = _project_with_segments(segments)

        orchestrator = GlobalPassOrchestrator()
        report = orchestrator.run(project)

        output_path = tmp_path / "output.docx"
        exporter = DocxExporter(project, global_report=report)
        result = exporter.export(output_path)

        assert result.exists()
        assert result.suffix == ".docx"

        doc = DocxDocument(str(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "chapter1" in full_text
        assert "chapter2" in full_text
        assert "本研究的引言" in full_text
        assert "结论与未来工作" in full_text

    def test_end_to_end_bilingual(self, tmp_path: Path) -> None:
        """Bilingual export includes source text alongside translations."""
        segments = [
            _seg("s1", "Introduction to the study.", "本研究的引言。", "chapter1"),
            _seg("s2", "Methods are described below.", "方法如下所述。", "chapter1"),
            _seg("s3", "Results show improvement.", "结果显示有所改善。", "chapter2"),
            _seg("s4", "Discussion of findings.", "研究发现的讨论。", "chapter2"),
            _seg("s5", "Conclusion and future work.", "结论与未来工作。", "chapter2"),
        ]
        project = _project_with_segments(segments)

        orchestrator = GlobalPassOrchestrator()
        report = orchestrator.run(project)

        output_path = tmp_path / "bilingual.docx"
        exporter = DocxExporter(project, bilingual=True, global_report=report)
        result = exporter.export(output_path)

        assert result.exists()

        doc = DocxDocument(str(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Introduction to the study." in full_text
        assert "本研究的引言" in full_text
        assert "Methods are described below." in full_text

    def test_global_pass_with_inconsistencies(self, tmp_path: Path) -> None:
        """Dropped citation causes report.passed=False; export still succeeds."""
        segments = [
            _seg(
                "s1",
                "As (Smith, 2020) demonstrated, results hold.",
                "结果成立。",
                "chapter1",
            ),
            _seg("s2", "Further analysis confirms.", "进一步分析证实。", "chapter1"),
        ]
        project = _project_with_segments(segments)

        orchestrator = GlobalPassOrchestrator()
        report = orchestrator.run(project)

        assert report.passed is False
        dropped = [c for c in report.citation_issues if c.issue_type == "dropped"]
        assert len(dropped) >= 1

        output_path = tmp_path / "with_issues.docx"
        exporter = DocxExporter(project, global_report=report)
        result = exporter.export(output_path)

        assert result.exists()

        doc = DocxDocument(str(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "ISSUES FOUND" in full_text
        assert "结果成立" in full_text
